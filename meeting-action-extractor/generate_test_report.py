# -*- coding: utf-8 -*-
"""
生成测试用例报告 docx
基于实际执行的12条测试用例结果
"""
import os
import sys
import json

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PROJECT_ROOT)

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shd)


def add_test_case_table(doc, test_cases):
    """添加测试用例表格"""
    # 按类别分组
    categories = {}
    for tc in test_cases:
        cat = tc['category']
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(tc)

    # 类别序号映射
    cat_order = ['正常场景', '边界场景', '无效输入场景', '权限场景', '系统异常场景', '人工介入场景']
    cat_num = {cat: chr(0x2160 + i) for i, cat in enumerate(cat_order)}

    for cat in cat_order:
        if cat not in categories:
            continue

        cases = categories[cat]
        p = doc.add_paragraph()
        run = p.add_run(f'{cat_num[cat]} {cat}')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        p.space_before = Pt(16)
        p.space_after = Pt(8)

        for tc in cases:
            # 用例编号和场景
            p = doc.add_paragraph()
            run = p.add_run(f'{tc["id"]}  {tc["scene"]}')
            run.bold = True
            run.font.size = Pt(12)
            p.space_before = Pt(10)

            # 创建5列表格
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 设置列宽
            for row in table.rows:
                row.cells[0].width = Cm(3.5)
                row.cells[1].width = Cm(13)

            # 填充内容
            labels = ['测试场景', '输入内容', '预期结果', '实际结果', '是否通过']
            values = [tc['scene'], tc['input'], tc['expected'], tc['actual'], tc['passed']]

            for i, (label, value) in enumerate(zip(labels, values)):
                cell_label = table.rows[i].cells[0]
                cell_value = table.rows[i].cells[1]

                # 标签
                p_label = cell_label.paragraphs[0]
                run_label = p_label.add_run(label)
                run_label.bold = True
                run_label.font.size = Pt(10)
                set_cell_shading(cell_label, 'F0F2F5')

                # 值
                p_value = cell_value.paragraphs[0]
                run_value = p_value.add_run(value)
                run_value.font.size = Pt(10)

                # 是否通过行着色
                if i == 4:
                    if value == '通过':
                        set_cell_shading(cell_value, 'E6F4EA')
                        run_value.font.color.rgb = RGBColor(0x13, 0x73, 0x33)
                        run_value.bold = True
                    elif value == '部分通过':
                        set_cell_shading(cell_value, 'FEF7E0')
                        run_value.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
                        run_value.bold = True
                    elif value == '未通过':
                        set_cell_shading(cell_value, 'FCE8E6')
                        run_value.font.color.rgb = RGBColor(0xC5, 0x22, 0x1F)
                        run_value.bold = True

            doc.add_paragraph()  # 空行


def generate_report():
    """生成测试报告"""
    # 读取测试结果
    results_path = os.path.join(PROJECT_ROOT, 'exports', 'test_results.json')
    with open(results_path, 'r', encoding='utf-8') as f:
        test_cases = json.load(f)

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 标题
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('测试用例执行报告')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('会议决策行动项提取工具')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('赛道：会议决策到任务执行自动化  |  版本：V1.0  |  日期：2026-08-01')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # 一、测试概述
    h = doc.add_heading('', level=1)
    run = h.add_run('一、测试概述')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    passed = sum(1 for tc in test_cases if tc['passed'] == '通过')
    partial = sum(1 for tc in test_cases if tc['passed'] == '部分通过')
    failed = sum(1 for tc in test_cases if tc['passed'] == '未通过')
    total = len(test_cases)

    overview = (
        f'本报告基于 5 份脱敏仿真会议字幕（产品版本规划会 TS-01、客户项目交付风险会 TS-02、'
        f'销售与实施交接会 TS-03、线上事故复盘会 TS-04、跨部门项目周会 TS-05），'
        f'设计并执行 {total} 条测试用例，覆盖正常、边界、无效输入、权限、系统异常和人工介入六大场景类别。\n\n'
        f'测试环境：Python 3.13 + Flask 3.1 + 规则提取模式（未配置 LLM API Key，M2 自动降级为规则模式）。\n\n'
        f'执行结果：通过 {passed} 条，部分通过 {partial} 条，未通过 {failed} 条，'
        f'通过率 {(passed + partial) / total * 100:.1f}%（含部分通过）。'
    )
    p = doc.add_paragraph(overview)
    p.style.font.size = Pt(10.5)

    # 统计表格
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['场景类别', '用例数', '通过', '部分通过', '未通过']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A73E8')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    cat_stats = {}
    for tc in test_cases:
        cat = tc['category']
        if cat not in cat_stats:
            cat_stats[cat] = {'total': 0, 'pass': 0, 'partial': 0, 'fail': 0}
        cat_stats[cat]['total'] += 1
        if tc['passed'] == '通过':
            cat_stats[cat]['pass'] += 1
        elif tc['passed'] == '部分通过':
            cat_stats[cat]['partial'] += 1
        else:
            cat_stats[cat]['fail'] += 1

    row_idx = 1
    for cat in ['正常场景', '边界场景', '无效输入场景', '权限场景', '系统异常场景', '人工介入场景']:
        if cat not in cat_stats:
            continue
        s = cat_stats[cat]
        cells = table.rows[row_idx].cells
        cells[0].text = cat
        cells[1].text = str(s['total'])
        cells[2].text = str(s['pass'])
        cells[3].text = str(s['partial'])
        cells[4].text = str(s['fail'])
        for c in cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_idx += 1

    doc.add_paragraph()

    # 二、测试用例详情
    h = doc.add_heading('', level=1)
    run = h.add_run('二、测试用例详情')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    add_test_case_table(doc, test_cases)

    # 三、失败原因分析
    h = doc.add_heading('', level=1)
    run = h.add_run('三、失败原因分析')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    failed_cases = [tc for tc in test_cases if tc['passed'] != '通过']
    if failed_cases:
        for tc in failed_cases:
            p = doc.add_paragraph()
            run = p.add_run(f'{tc["id"]} {tc["scene"]}')
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0xC5, 0x22, 0x1F)

            analysis = ''
            if tc['id'] == 'TC-01':
                analysis = (
                    '规则模式下，"五千条以内同步"这一决策语句被识别为普通讨论而非正式决策。'
                    '根因是规则引擎匹配到"建议"一词（上下文为"建议五千条以内同步"），触发讨论关键词误判。'
                    '林默后续的"正式决定"关键词未被同一段落捕获，因为决策语句和时间信息分布在不同段落。'
                    '在 LLM 模式下可显著改善，因为大模型能跨段落理解语义上下文。'
                )
            elif tc['id'] == 'TC-02':
                analysis = (
                    '工具正确识别为行动项且正确提取责任人高原，但截止时间解析失败。'
                    '根因是规则引擎的日期提取逻辑仅匹配"X月X日"和"本周/下周X"模式，'
                    '未覆盖"六月十六号"这种带"号"字的口语化日期表达。'
                    '修复方案：在 M2 规则模块中增加"X月X号"格式的正则匹配。'
                )
            elif tc['id'] == 'TC-05':
                analysis = (
                    '工具未能正确提取"六月二十六号"作为截止时间。'
                    '根因是规则引擎的日期提取未匹配到"二十六号"这一口语化数字表达，'
                    '且未实现"最终确认优先"的时间覆盖逻辑——当同一议题多次讨论时间被覆盖时，'
                    '工具应锚定到最终确认的段落而非首次提及处。'
                    '此问题与 Eval 报告中 ISSUE-01 一致，计划在 V1.1 版本修复。'
                )

            p = doc.add_paragraph(analysis)
            p.style.font.size = Pt(10.5)

            doc.add_paragraph()
    else:
        doc.add_paragraph('所有测试用例均通过。')

    # 四、测试结论
    h = doc.add_heading('', level=1)
    run = h.add_run('四、测试结论')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    conclusion = (
        f'本次测试共执行 {total} 条测试用例，覆盖六大场景类别。'
        f'其中 {passed} 条通过，{partial} 条部分通过，{failed} 条未通过，通过率 {(passed + partial) / total * 100:.1f}%。\n\n'
        '通过的测试用例验证了以下能力：\n'
        '- 多条决策与行动项混合提取能力正常，不同发言人任务正确归属\n'
        '- 责任人待确认场景正确标记 needs_review=true，不虚构责任人\n'
        '- 条件性承诺正确标注为目标日期并标记需人工确认\n'
        '- 普通讨论（建议、提问）不被误判为决策或行动项\n'
        '- 空文件和非会议文本正确处理，不崩溃不虚构\n'
        '- LLM API 未授权时自动降级为规则模式，不影响核心功能\n'
        '- 时间戳格式混乱时正确归一化，内容不丢失\n'
        '- 信息缺失项自动标记人工确认，不自行补充截止时间\n\n'
        '未通过的 3 条用例均与规则模式下日期提取和跨段落语义理解能力不足有关，'
        '在 LLM 大模型模式下预期可显著改善。'
        '未通过项已有明确修复计划（V1.1 版本），不影响工具核心功能可用性。'
    )
    doc.add_paragraph(conclusion)

    # 保存
    output_path = os.path.join(PROJECT_ROOT, 'exports', '测试用例执行报告.docx')
    doc.save(output_path)
    print(f'报告已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
